# app.py
import streamlit as st
import imaplib, email, os, io, base64
from email.header import decode_header
from email.message import EmailMessage
from PyPDF2 import PdfReader
from docx import Document

# Optional: Hugging Face summarizer (slow and heavy). Enable only if installed.
USE_HF = st.sidebar.checkbox("Use HuggingFace summarizer (slow/heavy)", value=False)
if USE_HF:
    from transformers import pipeline
    @st.cache_resource
    def load_summarizer():
        return pipeline("summarization", model="facebook/bart-large-cnn")
    summarizer = load_summarizer()
else:
    summarizer = None

st.set_page_config(page_title="IMAP Gmail → Contract Generator", layout="wide")
st.title("📧 Gmail (IMAP) → Contract Generator (App-password flow)")

st.markdown(
    """
**Important:** For Gmail you should create an **App password** (Google Account → Security → App passwords).
Do **not** enter your primary password unless you understand the risks.
"""
)

with st.sidebar.form("connect_form"):
    st.write("### Connect your Gmail (IMAP)")
    user_email = st.text_input("Gmail address", placeholder="you@gmail.com")
    user_pass = st.text_input("App password (or password)", type="password")
    mailbox = st.selectbox("Mailbox", ["INBOX", "ALL"], index=0)
    max_fetch = st.number_input("Emails to fetch", min_value=1, max_value=50, value=10)
    connect = st.form_submit_button("Connect")

imap_client = None
messages_list = []
if connect:
    if not user_email or not user_pass:
        st.error("Please enter email and app password.")
    else:
        try:
            st.info("Connecting to imap.gmail.com ...")
            imap_client = imaplib.IMAP4_SSL("imap.gmail.com")
            imap_client.login(user_email, user_pass)
            if mailbox == "INBOX":
                imap_client.select("INBOX")
            else:
                imap_client.select()  # default

            typ, data = imap_client.search(None, 'ALL')  # or 'UNSEEN'
            ids = data[0].split()[::-1]  # newest first
            ids = ids[:max_fetch]
            st.success(f"Connected. Found {len(ids)} messages. (showing latest {len(ids)})")

            # fetch basics
            for i, mid in enumerate(ids):
                typ, msg_data = imap_client.fetch(mid, "(RFC822)")
                raw = msg_data[0][1]
                msg = email.message_from_bytes(raw)
                subj, _ = decode_header(msg.get("Subject", ""))[0]
                if isinstance(subj, bytes):
                    subj = subj.decode(errors="ignore")
                from_ = msg.get("From", "")
                preview = msg.get("Date", "")  # minimal preview for dropdown key
                key = f"{i+1}: {subj} — {from_}"
                messages_list.append((key, mid, msg))
        except imaplib.IMAP4.error as e:
            st.error(f"IMAP error: {e}")
        except Exception as e:
            st.error(f"Connection failed: {e}")

# If connected and messages_list filled
if messages_list:
    keys = [m[0] for m in messages_list]
    selected = st.selectbox("Select an email to process", keys)
    sel_idx = keys.index(selected)
    _, sel_mid, sel_msg = messages_list[sel_idx]

    # Show subject/from/body snippet
    st.subheader("Email preview")
    subj = sel_msg.get("Subject", "")
    from_ = sel_msg.get("From", "")
    st.markdown(f"**Subject:** {subj}  \n**From:** {from_}")
    # get plain text body if available
    def get_body(message):
        if message.is_multipart():
            for part in message.walk():
                ctype = part.get_content_type()
                disp = str(part.get("Content-Disposition"))
                if ctype == "text/plain" and "attachment" not in disp:
                    return part.get_payload(decode=True).decode(errors="ignore")
        else:
            return message.get_payload(decode=True).decode(errors="ignore")
        return ""
    body_text = get_body(sel_msg) or "(no body text found)"
    st.text_area("Body (plain text)", body_text, height=200)

    # Download attachments (in-memory)
    attachments_text = ""
    attachments = []
    for part in sel_msg.walk():
        if part.get_content_maintype() == "multipart":
            continue
        if part.get("Content-Disposition") is None:
            continue
        filename = part.get_filename()
        if filename:
            payload = part.get_payload(decode=True)
            # save to temp file in memory
            attachments.append((filename, payload))
            st.write(f"Attachment found: {filename}")

            # Extract text from common types
            if filename.lower().endswith(".pdf"):
                try:
                    reader = PdfReader(io.BytesIO(payload))
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            attachments_text += text + "\n"
                except Exception as e:
                    st.warning(f"Could not extract PDF text from {filename}: {e}")
            elif filename.lower().endswith(".docx"):
                try:
                    # write to temp on disk then read via python-docx
                    with open("/tmp/_temp.docx", "wb") as tf:
                        tf.write(payload)
                    doc = Document("/tmp/_temp.docx")
                    for p in doc.paragraphs:
                        attachments_text += p.text + "\n"
                    os.remove("/tmp/_temp.docx")
                except Exception as e:
                    st.warning(f"Could not extract DOCX text from {filename}: {e}")
            elif filename.lower().endswith(".txt"):
                try:
                    attachments_text += payload.decode(errors="ignore") + "\n"
                except:
                    pass
            else:
                st.info(f"Unsupported attachment type for text extraction: {filename}")

    st.subheader("Extracted text from attachments")
    st.text_area("Attachments extracted text", attachments_text, height=200)

    # Summarize
    def simple_summary(text, max_chars=800):
        return text[:max_chars] + ("..." if len(text) > max_chars else "")

    if st.button("Summarize & Generate contract"):
        combined_text = body_text + "\n\n" + attachments_text
        if not combined_text.strip():
            st.error("No content to summarize from this email/attachments.")
        else:
            if summarizer:
                # HF model expects shorter inputs; for long docs you should chunk
                try:
                    s = summarizer(combined_text, max_length=150, min_length=40, do_sample=False)
                    summary = s[0]["summary_text"]
                except Exception as e:
                    st.warning(f"HuggingFace summarizer failed: {e}. Falling back to simple summary.")
                    summary = simple_summary(combined_text)
            else:
                summary = simple_summary(combined_text)

            # Create a DOCX contract
            contract = Document()
            contract.add_heading("Generated Contract", level=0)
            contract.add_paragraph("Based on the selected email and attachments.")
            contract.add_heading("Summary", level=1)
            contract.add_paragraph(summary)
            contract.add_heading("Original Email Body", level=1)
            contract.add_paragraph(body_text)
            if attachments_text.strip():
                contract.add_heading("Attachments (extracted)", level=1)
                contract.add_paragraph(attachments_text)

            out_name = "generated_contract.docx"
            contract.save(out_name)

            # Provide download
            with open(out_name, "rb") as f:
                data = f.read()
            b64 = base64.b64encode(data).decode()
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="{out_name}">⬇️ Download contract.docx</a>'
            st.markdown(href, unsafe_allow_html=True)

            st.success("Contract generated. Click download link above.")
